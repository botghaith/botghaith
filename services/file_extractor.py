import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def extract_text_from_file(file_path: Path) -> str:
    suffix = file_path.suffix.lower()

    if suffix == ".txt":
        return _extract_txt(file_path)
    elif suffix == ".pdf":
        return _extract_pdf(file_path)
    elif suffix in (".docx", ".doc"):
        return _extract_docx(file_path)
    else:
        raise ValueError(f"نوع الملف غير مدعوم: {suffix}")


def _extract_txt(path: Path) -> str:
    for enc in ("utf-8", "utf-16", "cp1256", "latin-1"):
        try:
            return path.read_text(encoding=enc)
        except (UnicodeDecodeError, UnicodeError):
            continue
    return path.read_text(encoding="utf-8", errors="replace")


def _extract_pdf(path: Path) -> str:
    import fitz

    doc = fitz.open(str(path))
    text_parts = [page.get_text() for page in doc]
    doc.close()
    native = "\n".join(text_parts).strip()
    compact = native.replace(" ", "").replace("\n", "")
    if len(compact) >= 20:
        return native

    try:
        from services.ocr_service import extract_pdf_text_smart

        ocr = extract_pdf_text_smart(path, path.parent).strip()
        if ocr:
            logger.info("PDF OCR fallback extracted %d chars", len(ocr))
            return ocr
    except Exception as e:
        logger.warning("PDF OCR fallback failed: %s", e)

    return native


def _extract_docx(path: Path) -> str:
    from docx import Document
    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)
    return "\n".join(paragraphs)
