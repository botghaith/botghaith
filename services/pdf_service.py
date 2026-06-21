import logging
from pathlib import Path

from services.text_shape import has_arabic, shape_for_pdf

logger = logging.getLogger(__name__)


def images_to_pdf(image_paths: list[Path], output_path: Path):
    import img2pdf
    with open(output_path, "wb") as f:
        f.write(img2pdf.convert([str(p) for p in image_paths]))


def pdf_to_images(pdf_path: Path, output_dir: Path) -> list[Path]:
    import fitz
    doc = fitz.open(str(pdf_path))
    images = []
    for i, page in enumerate(doc):
        pix = page.get_pixmap(dpi=150)
        img_path = output_dir / f"page_{i + 1}.png"
        pix.save(str(img_path))
        images.append(img_path)
    doc.close()
    return images


def word_to_pdf(docx_path: Path, output_path: Path):
    from docx import Document
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    doc = Document(str(docx_path))
    c = canvas.Canvas(str(output_path), pagesize=A4)
    width, height = A4
    y = height - 50

    _set_font(c)
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            y -= 10
            continue
        lines = _wrap_text(text, 80)
        for line in lines:
            if y < 50:
                c.showPage()
                y = height - 50
                _set_font(c)
            c.drawString(50, y, line)
            y -= 18
    c.save()


def pdf_to_word(pdf_path: Path, output_path: Path):
    import fitz
    from docx import Document

    doc = fitz.open(str(pdf_path))
    word_doc = Document()
    for page in doc:
        text = page.get_text()
        if text.strip():
            word_doc.add_paragraph(text)
    doc.close()
    word_doc.save(str(output_path))


def merge_pdfs(pdf_paths: list[Path], output_path: Path):
    import fitz
    merged = fitz.open()
    for p in pdf_paths:
        src = fitz.open(str(p))
        merged.insert_pdf(src)
        src.close()
    merged.save(str(output_path))
    merged.close()


def split_pdf(pdf_path: Path, output_dir: Path) -> list[Path]:
    import fitz
    doc = fitz.open(str(pdf_path))
    files = []
    for i in range(len(doc)):
        new_doc = fitz.open()
        new_doc.insert_pdf(doc, from_page=i, to_page=i)
        out = output_dir / f"part_{i + 1}.pdf"
        new_doc.save(str(out))
        new_doc.close()
        files.append(out)
    doc.close()
    return files


def compress_pdf(pdf_path: Path, output_path: Path):
    import fitz
    doc = fitz.open(str(pdf_path))
    doc.save(str(output_path), garbage=4, deflate=True, clean=True)
    doc.close()


def extract_pdf_text(pdf_path: Path) -> str:
    from services.ocr_service import extract_pdf_text_smart
    return extract_pdf_text_smart(pdf_path, pdf_path.parent)


def reorder_pdf_pages(pdf_path: Path, order: list[int], output_path: Path):
    import fitz
    doc = fitz.open(str(pdf_path))
    new_doc = fitz.open()
    for page_num in order:
        if 0 <= page_num < len(doc):
            new_doc.insert_pdf(doc, from_page=page_num, to_page=page_num)
    new_doc.save(str(output_path))
    new_doc.close()
    doc.close()


def create_text_pdf(text: str, output_path: Path, title: str = ""):
    create_bilingual_pdf(text, output_path, title=title)


def create_bilingual_pdf(text: str, output_path: Path, title: str = ""):
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    c = canvas.Canvas(str(output_path), pagesize=A4)
    width, height = A4
    y = height - 50
    font_name = _set_font(c)

    if title:
        c.setFont(font_name, 16)
        for line in _pdf_lines(title, 70):
            if y < 60:
                c.showPage()
                y = height - 50
                _set_font(c)
            c.drawRightString(width - 50, y, line) if has_arabic(line) else c.drawString(50, y, line)
            y -= 22
        y -= 10
        _set_font(c)

    for line in text.split("\n"):
        if y < 60:
            c.showPage()
            y = height - 50
            _set_font(c)
        if not line.strip():
            y -= 10
            continue
        for w in _pdf_lines(line, 85):
            if y < 60:
                c.showPage()
                y = height - 50
                _set_font(c)
            shaped = shape_for_pdf(w)
            if has_arabic(shaped):
                c.drawRightString(width - 50, y, shaped)
            else:
                c.drawString(50, y, shaped)
            y -= 16
        y -= 2

    c.save()


def _ensure_page(c, y, height, font_name):
    if y < 70:
        c.showPage()
        _set_font(c)
        return height - 50
    return y


def _draw_wrapped(c, width, y, height, text: str, font_name: str, size: int, rtl: bool, margin: int = 55):
    c.setFont(font_name, size)
    for part in _pdf_lines(text, 78):
        y = _ensure_page(c, y, height, font_name)
        c.setFont(font_name, size)
        shaped = shape_for_pdf(part) if has_arabic(part) else part
        if rtl:
            c.drawRightString(width - margin, y, shaped)
        else:
            c.drawString(margin, y, shaped)
        y -= size + 5
    return y


def create_pairs_pdf(
    pairs: list[tuple[str, str]],
    output_path: Path,
    title: str = "",
    direction: str = "en_ar",
):
    """PDF مرتب: كل فقرة/سطر وأصله وترجمته بشكل منظم"""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.pdfgen import canvas

    c = canvas.Canvas(str(output_path), pagesize=A4)
    width, height = A4
    y = height - 45
    font_name = _set_font(c)

    if title:
        y = _draw_wrapped(c, width, y, height, title, font_name, 15, has_arabic(title))
        y -= 8
        c.setStrokeColor(colors.grey)
        c.line(55, y, width - 55, y)
        y -= 18

    src_rtl = direction == "ar_en"
    tr_rtl = direction == "en_ar"

    for i, (src, tr) in enumerate(pairs, 1):
        y = _ensure_page(c, y, height, font_name)
        c.setFont(font_name, 9)
        c.setFillColor(colors.HexColor("#555555"))
        label = shape_for_pdf(f"— {i} —") if has_arabic(f"— {i} —") else f"— {i} —"
        c.drawString(55, y, label)
        y -= 16
        c.setFillColor(colors.black)

        c.setFont(font_name, 10)
        c.drawString(55, y, "Original / الأصل:")
        y -= 14
        y = _draw_wrapped(c, width, y, height, src, font_name, 12, src_rtl)

        y -= 4
        c.setFont(font_name, 10)
        c.drawString(55, y, "Translation / الترجمة:")
        y -= 14
        y = _draw_wrapped(c, width, y, height, tr, font_name, 12, tr_rtl)

        y -= 6
        c.setStrokeColor(colors.HexColor("#CCCCCC"))
        c.line(55, y, width - 55, y)
        y -= 14

    c.save()


def create_literal_pdf(
    sections: list[tuple[str, list[tuple[str, str]]]],
    output_path: Path,
    title: str = "",
    direction: str = "en_ar",
):
    """PDF حرفي مرتب حسب الفقرات"""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.pdfgen import canvas

    c = canvas.Canvas(str(output_path), pagesize=A4)
    width, height = A4
    y = height - 45
    font_name = _set_font(c)

    if title:
        y = _draw_wrapped(c, width, y, height, title, font_name, 15, has_arabic(title))
        y -= 12

    for sec_title, pairs in sections:
        y = _ensure_page(c, y, height, font_name)
        c.setFont(font_name, 11)
        c.setFillColor(colors.HexColor("#1a5276"))
        sec = shape_for_pdf(sec_title) if has_arabic(sec_title) else sec_title
        c.drawString(55, y, sec)
        y -= 14
        c.setFillColor(colors.black)
        c.setStrokeColor(colors.HexColor("#1a5276"))
        c.line(55, y + 10, width - 55, y + 10)
        y -= 6

        for word, tr in pairs:
            y = _ensure_page(c, y, height, font_name)
            line = f"{word}  —  {tr}"
            rtl = has_arabic(tr) or has_arabic(word)
            y = _draw_wrapped(c, width, y, height, line, font_name, 11, rtl)
            y -= 2
        y -= 10

    c.save()


def _pdf_lines(text: str, max_chars: int) -> list[str]:
    return _wrap_text(text, max_chars) if text else [""]


def _set_font(c):
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    font_paths = [
        "C:/Windows/Fonts/tahoma.ttf",
        "C:/Windows/Fonts/segoeui.ttf",
        "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/trado.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ]
    for fp in font_paths:
        try:
            if Path(fp).exists():
                name = "TrFont"
                if name not in pdfmetrics.getRegisteredFontNames():
                    pdfmetrics.registerFont(TTFont(name, fp))
                c.setFont(name, 11)
                return name
        except Exception:
            continue
    c.setFont("Helvetica", 11)
    return "Helvetica"


def _wrap_text(text: str, max_chars: int) -> list[str]:
    if len(text) <= max_chars:
        return [text]
    words = text.split()
    lines = []
    current = ""
    for w in words:
        if len(current) + len(w) + 1 <= max_chars:
            current = f"{current} {w}".strip()
        else:
            if current:
                lines.append(current)
            current = w
    if current:
        lines.append(current)
    return lines if lines else [text[:max_chars]]
