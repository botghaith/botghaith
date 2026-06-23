"""
ترجمة الملفات بأربعة أنماط:
1) حرفي — كل كلمة بسطر وحدها مع ترجمتها
2) بنفس الترتيب — نفس هيكل الملف مترجماً بالكامل
3) سطر بسطر — كل فقرة وترجمتها تحتها
4) فوق الكلمات — نفس الملف مع ترجمة صغيرة فوق كل كلمة
"""
import logging
import os
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.run import Run

from services.file_extractor import extract_text_from_file
from services.pdf_service import create_bilingual_pdf, create_pairs_pdf, create_literal_pdf
from services.translator import translate_text, resolve_direction, set_file_translation_mode, is_translator_ready
from config import use_fast_file_translation, prefer_local_for_files, is_render_host, file_max_paragraphs, use_dual_file_translation, use_full_file_translation
from services.text_shape import (
    is_mostly_arabic,
    find_arabic_font,
    shape_for_pdf,
    set_run_font,
    set_paragraph_direction,
    style_paragraph,
)

logger = logging.getLogger(__name__)

MAX_CHARS = 500_000
MAX_FILE_MB = 50
WORD_TOKEN_RE = re.compile(r"\S+")
WORD_CHAR_RE = re.compile(r"[\w\u0600-\u06FF]", re.UNICODE)
_word_cache: dict[tuple[str, str], str] = {}
_WORD_BATCH_SIZE = 40


def _clear_word_cache():
    _word_cache.clear()


def _prewarm_word_cache_for_text(text: str, direction: str) -> None:
    """ترجمة الكلمات الفريدة دفعة واحدة — أسرع بكثير من كلمة/طلب."""
    direction = resolve_direction(text, direction)
    unique: list[str] = []
    seen: set[str] = set()
    for unit in _extract_logical_units(text):
        for token in WORD_TOKEN_RE.findall(unit):
            key = token.lower()
            if key not in seen and WORD_CHAR_RE.search(token):
                seen.add(key)
                unique.append(token)
    if not unique:
        return

    logger.info("Prewarming word cache: %d unique tokens", len(unique))
    for i in range(0, len(unique), _WORD_BATCH_SIZE):
        chunk = unique[i:i + _WORD_BATCH_SIZE]
        payload = "\n".join(f"{idx + 1}. {w}" for idx, w in enumerate(chunk))
        try:
            translated = translate_text(payload, direction)
            lines = [ln.strip() for ln in translated.splitlines() if ln.strip()]
            parsed: list[str] = []
            for ln in lines:
                m = re.match(r"^\d+\.\s*(.+)$", ln)
                parsed.append(m.group(1).strip() if m else ln)
            if len(parsed) == len(chunk):
                for w, tw in zip(chunk, parsed):
                    _word_cache[(w.lower(), direction)] = tw
                continue
        except Exception as e:
            logger.warning("Batch word prewarm failed: %s", e)
        for w in chunk:
            key = (w.lower(), direction)
            if key not in _word_cache:
                try:
                    _word_cache[key] = translate_text(w, direction)
                except Exception:
                    _word_cache[key] = w


def translate_word(word: str, direction: str) -> str:
    raw = word.strip()
    if not raw or not WORD_CHAR_RE.search(raw):
        return word
    core = re.sub(r"^[^\w\u0600-\u06FF]+|[^\w\u0600-\u06FF]+$", "", raw)
    if not core:
        return word
    prefix = raw[: raw.index(core)] if core in raw else ""
    suffix = raw[raw.index(core) + len(core) :] if core in raw else ""
    key = (core.lower(), direction)
    if key not in _word_cache:
        _word_cache[key] = translate_text(core, direction)
    return f"{prefix}{_word_cache[key]}{suffix}"


def _check_file_limits(source_path: Path, content: str):
    size_mb = source_path.stat().st_size / (1024 * 1024)
    if size_mb > MAX_FILE_MB:
        raise ValueError(f"الملف كبير ({size_mb:.1f} MB). الحد الأقصى {MAX_FILE_MB} MB")
    if len(content) > MAX_CHARS:
        raise ValueError(f"النص طويل جداً ({len(content)} حرف). الحد الأقصى {MAX_CHARS} حرف")


def _extract_words(text: str) -> list[str]:
    return [t for t in WORD_TOKEN_RE.findall(text) if WORD_CHAR_RE.search(t)]


def _apply_affixes(original: str, translated_core: str) -> str:
    core = re.sub(r"^[^\w\u0600-\u06FF]+|[^\w\u0600-\u06FF]+$", "", original)
    if not core or core not in original:
        return translated_core
    prefix = original[: original.index(core)]
    suffix = original[original.index(core) + len(core) :]
    return f"{prefix}{translated_core}{suffix}"


def translate_word_in_context(token: str, line: str, direction: str) -> str:
    """ترجمة أدق: تستخدم سياق الجملة عند تطابق عدد الكلمات"""
    words = [w for w in WORD_TOKEN_RE.findall(line) if WORD_CHAR_RE.search(w)]
    if token not in words:
        return translate_word(token, direction)

    idx = words.index(token)
    for window in (1, 2, 3):
        start = max(0, idx - window + 1)
        end = min(len(words), idx + window)
        chunk_words = words[start:end]
        chunk_tr = translate_text(" ".join(chunk_words), direction).strip()
        tr_tokens = [w for w in WORD_TOKEN_RE.findall(chunk_tr) if WORD_CHAR_RE.search(w)]
        if len(tr_tokens) == len(chunk_words):
            rel = idx - start
            return _apply_affixes(token, tr_tokens[rel])
    return translate_word(token, direction)


def _extract_logical_units(text: str) -> list[str]:
    """دمج الأسطر المتقطعة في فقرات/جمل كاملة لترجمة أدق"""
    units: list[str] = []
    buffer: list[str] = []

    for line in text.splitlines():
        s = line.strip()
        if not s:
            if buffer:
                units.append(" ".join(buffer))
                buffer = []
            continue
        buffer.append(s)
        if re.search(r"[.!?؟…:;]$", s) or len(" ".join(buffer)) > 280:
            units.append(" ".join(buffer))
            buffer = []

    if buffer:
        units.append(" ".join(buffer))
    return units if units else [text.strip()]


def _format_word_pair(token: str, translation: str) -> str:
    return f"{token}  —  {translation}"


def build_literal_sections(text: str, direction: str) -> list[tuple[str, list[tuple[str, str]]]]:
    """ترجمة حرفية مرتبة حسب الفقرات — الملف كاملاً"""
    direction = resolve_direction(text, direction)
    sections: list[tuple[str, list[tuple[str, str]]]] = []

    for idx, unit in enumerate(_extract_logical_units(text), 1):
        pairs: list[tuple[str, str]] = []
        for token in WORD_TOKEN_RE.findall(unit):
            if not WORD_CHAR_RE.search(token):
                continue
            pairs.append((token, translate_word_in_context(token, unit, direction)))
        if pairs:
            sections.append((f"الفقرة {idx}", pairs))

    return sections


def build_literal_sections_from_pairs(
    pairs: list[tuple[str, str]],
) -> list[tuple[str, list[tuple[str, str]]]]:
    """حرفي من فقرات مترجمة — بديل سريع عند غياب Argos."""
    sections: list[tuple[str, list[tuple[str, str]]]] = []
    for idx, (unit, tr) in enumerate(pairs, 1):
        src_words = [t for t in WORD_TOKEN_RE.findall(unit) if WORD_CHAR_RE.search(t)]
        tr_words = [t for t in WORD_TOKEN_RE.findall(tr) if WORD_CHAR_RE.search(t)]
        if src_words and len(src_words) == len(tr_words):
            word_pairs = list(zip(src_words, tr_words))
        elif src_words:
            word_pairs = [(w, tr) for w in src_words]
        else:
            word_pairs = [(unit, tr)]
        sections.append((f"الفقرة {idx}", word_pairs))
    return sections


def _build_literal_files_from_pairs(
    pairs: list[tuple[str, str]], direction: str, out_dir: Path, stem: str,
) -> dict[str, Path]:
    sections = build_literal_sections_from_pairs(pairs)
    literal_pdf = out_dir / f"{stem}_1_حرفي.pdf"
    create_literal_pdf(sections, literal_pdf, title="ترجمة حرفية — كلمة بكلمة", direction=direction)
    return {"literal": literal_pdf}


def build_literal_text(text: str, direction: str) -> str:
    """كل كلمة مع ترجمتها في سطر واحد — مرتب حسب الفقرات"""
    lines: list[str] = []
    for title, pairs in build_literal_sections(text, direction):
        lines.append(f"{'═' * 12} {title} {'═' * 12}")
        for word, tr in pairs:
            lines.append(_format_word_pair(word, tr))
        lines.append("")
    return "\n".join(lines).strip()


def _build_literal_files(content: str, direction: str, out_dir: Path, stem: str) -> dict[str, Path]:
    direction = resolve_direction(content, direction)
    sections = build_literal_sections(content, direction)

    literal_pdf = out_dir / f"{stem}_1_حرفي.pdf"
    create_literal_pdf(sections, literal_pdf, title="ترجمة حرفية — كلمة بكلمة", direction=direction)

    return {"literal": literal_pdf}


def build_line_pairs(text: str, direction: str) -> list[tuple[str, str]]:
    """ترجمة دقيقة: فقرة كاملة ثم ترجمتها"""
    direction = resolve_direction(text, direction)
    pairs: list[tuple[str, str]] = []

    for unit in _extract_logical_units(text):
        unit = unit.strip()
        if not unit:
            continue
        # ترجمة الجملة/الفقرة كاملة لدقة أعلى
        translated = translate_text(unit, direction)
        pairs.append((unit, translated))

    return pairs


def _build_line_pairs_file(
    content: str,
    direction: str,
    out_dir: Path,
    stem: str,
    pairs: list[tuple[str, str]] | None = None,
) -> dict[str, Path]:
    direction = resolve_direction(content, direction)
    if pairs is None:
        pairs = build_line_pairs(content, direction)
    path = out_dir / f"{stem}_3_سطر_بسطر.pdf"
    create_pairs_pdf(pairs, path, title="ترجمة سطر بسطر — أصل وترجمة", direction=direction)
    return {"line_pairs": path}


def _write_structured_online(
    source_path: Path,
    pairs: list[tuple[str, str]],
    out_dir: Path,
    stem: str,
    direction: str,
) -> Path:
    suffix = source_path.suffix.lower()
    full_text = "\n\n".join(tr for _, tr in pairs)

    if suffix in (".docx", ".doc"):
        path = out_dir / f"{stem}_2_بنفس_الترتيب.docx"
        doc = Document()
        for _, tr in pairs:
            para = doc.add_paragraph(tr)
            style_paragraph(para, tr, direction)
        doc.save(path)
        return path

    if suffix == ".pdf":
        path = out_dir / f"{stem}_2_بنفس_الترتيب.pdf"
        create_pairs_pdf(
            pairs, path,
            title="ترجمة بنفس الترتيب — أصل وترجمة",
            direction=direction,
        )
        return path

    path = out_dir / f"{stem}_2_بنفس_الترتيب.txt"
    path.write_text(full_text, encoding="utf-8-sig")
    return path


def _build_overlay_online(
    pairs: list[tuple[str, str]], out_dir: Path, stem: str, direction: str,
) -> dict[str, Path]:
    """ملف 4 — كلمة وترجمتها (من محاذاة الفقرات)."""
    sections = build_literal_sections_from_pairs(pairs)
    flat_pairs: list[tuple[str, str]] = []
    for _, word_pairs in sections:
        flat_pairs.extend(word_pairs)
    path = out_dir / f"{stem}_4_فوق_الكلمات.pdf"
    create_pairs_pdf(
        [(_format_word_pair(w, t), "") for w, t in flat_pairs],
        path,
        title="ترجمة فوق الكلمات",
        direction=direction,
    )
    return {"overlay": path}


def _paragraph_pairs(content: str, direction: str) -> list[tuple[str, str]]:
    direction = resolve_direction(content, direction)
    pairs: list[tuple[str, str]] = []
    units = _extract_logical_units(content)
    total = len(units)
    for i, unit in enumerate(units, 1):
        unit = unit.strip()
        if not unit:
            continue
        if i % 5 == 0 or i == total:
            logger.info("Translating paragraph %d/%d", i, total)
        pairs.append((unit, translate_text(unit, direction)))
    return pairs


def prepare_online_file_translation(source_path: Path, direction: str = "auto") -> dict:
    """استخراج + ترجمة فقرات — الخطوة البطيئة."""
    content = extract_text_from_file(source_path)
    if not content.strip():
        raise ValueError("لم يتم العثور على نص في الملف")
    _check_file_limits(source_path, content)
    direction = resolve_direction(content, direction)
    pairs = _paragraph_pairs(content, direction)
    if not pairs:
        raise ValueError("لم يتم العثور على فقرات للترجمة")

    truncated = False
    limit = file_max_paragraphs()
    if len(pairs) > limit:
        pairs = pairs[:limit]
        truncated = True
        logger.warning("Truncated to %d paragraphs", limit)

    return {
        "content": content,
        "direction": direction,
        "pairs": pairs,
        "stem": source_path.stem,
        "source_path": source_path,
        "truncated": truncated,
    }


def prepare_online_image_translation(image_path: Path, direction: str = "auto") -> dict:
    from services.ocr_service import ocr_image

    content = ocr_image(image_path)
    if not content.strip():
        raise ValueError("لم يتم العثور على نص في الصورة")
    _check_file_limits(image_path, content)
    direction = resolve_direction(content, direction)
    pairs = _paragraph_pairs(content, direction)
    if not pairs:
        raise ValueError("لم يتم العثور على نص للترجمة في الصورة")

    truncated = False
    limit = file_max_paragraphs()
    if len(pairs) > limit:
        pairs = pairs[:limit]
        truncated = True

    return {
        "content": content,
        "direction": direction,
        "pairs": pairs,
        "stem": image_path.stem,
        "source_path": image_path,
        "truncated": truncated,
        "is_image": True,
    }


def build_online_literal(data: dict, output_dir: Path) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    return _build_literal_files_from_pairs(
        data["pairs"], data["direction"], output_dir, data["stem"],
    )["literal"]


def build_online_structured(data: dict, output_dir: Path) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    if data.get("is_image"):
        path = output_dir / f"{data['stem']}_2_بنفس_الترتيب.pdf"
        create_pairs_pdf(
            data["pairs"], path,
            title="ترجمة الصورة — أصل وترجمة",
            direction=data["direction"],
        )
        return path
    return _write_structured_online(
        data["source_path"], data["pairs"], output_dir, data["stem"], data["direction"],
    )


def build_online_line_pairs(data: dict, output_dir: Path) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    return _build_line_pairs_file(
        data["content"], data["direction"], output_dir, data["stem"], pairs=data["pairs"],
    )["line_pairs"]


def build_online_overlay(data: dict, output_dir: Path) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    return _build_overlay_online(
        data["pairs"], output_dir, data["stem"], data["direction"],
    )["overlay"]


def _argos_available() -> bool:
    return prefer_local_for_files() and is_translator_ready()


def _translate_file_online_full(
    source_path: Path, output_dir: Path, direction: str = "auto",
) -> dict[str, Path]:
    """4 ملفات عبر الإنternet — فقرة بفقرة (لا تعليق)."""
    _clear_word_cache()
    stem = source_path.stem
    output_dir.mkdir(parents=True, exist_ok=True)

    content = extract_text_from_file(source_path)
    if not content.strip():
        raise ValueError("لم يتم العثور على نص في الملف")
    _check_file_limits(source_path, content)
    direction = resolve_direction(content, direction)

    logger.info("Online 4-file translation: %s (%d chars)", source_path.name, len(content))
    pairs = _paragraph_pairs(content, direction)
    if not pairs:
        raise ValueError("لم يتم العثور على فقرات للترجمة")

    result = _build_literal_files_from_pairs(pairs, direction, output_dir, stem)
    result["structured"] = _write_structured_online(source_path, pairs, output_dir, stem, direction)
    result.update(_build_line_pairs_file(content, direction, output_dir, stem, pairs=pairs))
    result.update(_build_overlay_online(pairs, output_dir, stem, direction))
    logger.info("Online 4-file translation done")
    return result


def _translate_image_online_full(
    image_path: Path, output_dir: Path, direction: str = "auto",
) -> dict[str, Path]:
    from services.ocr_service import ocr_image

    _clear_word_cache()
    stem = image_path.stem
    output_dir.mkdir(parents=True, exist_ok=True)

    content = ocr_image(image_path)
    if not content.strip():
        raise ValueError("لم يتم العثور على نص في الصورة")
    _check_file_limits(image_path, content)
    direction = resolve_direction(content, direction)

    logger.info("Online 4-file image translation: %s", image_path.name)
    pairs = _paragraph_pairs(content, direction)
    if not pairs:
        raise ValueError("لم يتم العثور على نص للترجمة في الصورة")

    result = _build_literal_files_from_pairs(pairs, direction, output_dir, stem)
    structured_pdf = output_dir / f"{stem}_2_بنفس_الترتيب.pdf"
    create_pairs_pdf(pairs, structured_pdf, title="ترجمة الصورة — أصل وترجمة", direction=direction)
    result["structured"] = structured_pdf
    result.update(_build_line_pairs_file(content, direction, output_dir, stem, pairs=pairs))
    result.update(_build_overlay_online(pairs, output_dir, stem, direction))
    return result


OVERLAY_TR_SIZE = 6
OVERLAY_TR_SIZE_PDF = 7
OVERLAY_WORD_SIZE = 11
OVERLAY_LINE_SPACING = 0.85


def _paragraph_has_image(para) -> bool:
    for run in para.runs:
        if run._element.xpath(".//w:drawing") or run._element.xpath(".//w:pict"):
            return True
    return False


def _remove_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tbl_pr.append(borders)


def _set_cell_tight_margin(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for margin in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{margin}")
        el.set(qn("w:w"), "0")
        el.set(qn("w:type"), "dxa")
        tc_mar.append(el)
    tc_pr.append(tc_mar)


def _set_cell_valign(cell, align: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    v_align = OxmlElement("w:vAlign")
    v_align.set(qn("w:val"), align)
    tc_pr.append(v_align)


def _set_table_tight_spacing(table):
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tbl_pr)
    spacing = OxmlElement("w:tblCellSpacing")
    spacing.set(qn("w:w"), "0")
    spacing.set(qn("w:type"), "dxa")
    tbl_pr.append(spacing)


def _iter_pdf_word_boxes(page):
    for block in page.get_text("dict").get("blocks", []):
        if block.get("type") != 0:
            continue
        for line in block.get("lines", []):
            line_text = "".join(s.get("text", "") for s in line.get("spans", []))
            for span in line.get("spans", []):
                span_text = span.get("text", "")
                if not span_text.strip():
                    continue
                x0, y0, x1, y1 = span["bbox"]
                span_w = max(x1 - x0, 1)
                total = max(len(span_text), 1)
                for match in WORD_TOKEN_RE.finditer(span_text):
                    token = match.group()
                    if not WORD_CHAR_RE.search(token):
                        continue
                    start, end = match.start(), match.end()
                    wx0 = x0 + (start / total) * span_w
                    wx1 = x0 + (end / total) * span_w
                    yield token, wx0, y0, wx1, y1, line_text


def _pdf_insert_translation_above(page, x0, y0, x1, y1, text: str, fontfile: str | None, rtl: bool):
    import fitz
    from services.text_shape import has_arabic

    display = shape_for_pdf(text) if has_arabic(text) else text
    word_w = max(x1 - x0, 3)
    word_h = max(y1 - y0, 3)
    fs = min(OVERLAY_TR_SIZE_PDF, word_h * 0.42)
    fs = max(fs, 5.5)

    font_kwargs = {}
    if fontfile:
        font_kwargs["fontfile"] = fontfile
        font_kwargs["fontname"] = "TahomaAr"

    def _text_width(size: float) -> float:
        try:
            return fitz.get_text_length(display, fontsize=size, **font_kwargs)
        except Exception:
            return len(display) * size * 0.45

    tw = _text_width(fs)
    while tw > word_w * 1.05 and fs > 4.0:
        fs -= 0.15
        tw = _text_width(fs)

    x = max(x0, x1 - tw) if rtl else x0
    # baseline قريب جداً من أعلى الكلمة — مسافة ضئيلة بين الترجمة والكلمة
    y = y0 + fs * 0.12

    try:
        page.insert_text((x, y), display, fontsize=fs, **font_kwargs)
    except Exception:
        page.insert_text((x, y), display, fontsize=fs)


def _run_has_image(run) -> bool:
    return bool(run._element.xpath(".//w:drawing") or run._element.xpath(".//w:pict"))


def _add_overlay_runs_at_index(
    para, parent, insert_idx: int, text: str, direction: str
) -> int:
    direction = resolve_direction(text, direction)
    for token in WORD_TOKEN_RE.findall(text):
        if WORD_CHAR_RE.search(token):
            tr = translate_word_in_context(token, text, direction)
            r = OxmlElement("w:r")
            parent.insert(insert_idx, r)
            insert_idx += 1
            tr_run = Run(r, para)
            tr_run.text = tr
            set_run_font(tr_run, "Tahoma", OVERLAY_TR_SIZE)
            tr_run.font.superscript = True
            tr_run.font.size = Pt(OVERLAY_TR_SIZE)
            r = OxmlElement("w:r")
            parent.insert(insert_idx, r)
            insert_idx += 1
            w_run = Run(r, para)
            w_run.text = token
            set_run_font(w_run, "Tahoma", OVERLAY_WORD_SIZE)
            if not token.endswith((" ", "\t")):
                sp = OxmlElement("w:r")
                parent.insert(insert_idx, sp)
                insert_idx += 1
                Run(sp, para).text = " "
        else:
            r = OxmlElement("w:r")
            parent.insert(insert_idx, r)
            insert_idx += 1
            Run(r, para).text = token
    return insert_idx


def _transform_text_run(para, run, direction: str):
    if _run_has_image(run):
        return
    segment = run.text
    if not segment.strip():
        return
    elem = run._element
    parent = elem.getparent()
    idx = parent.index(elem)
    parent.remove(elem)
    _add_overlay_runs_at_index(para, parent, idx, segment, direction)


def _set_row_exact_height(row, height_pt: float):
    tr_pr = row._tr.get_or_add_trPr()
    tr_height = OxmlElement("w:trHeight")
    tr_height.set(qn("w:val"), str(int(height_pt * 20)))
    tr_height.set(qn("w:hRule"), "exact")
    tr_pr.append(tr_height)


def _replace_paragraph_with_overlay_table(doc: Document, para, direction: str):
    text = para.text.strip()
    tokens = WORD_TOKEN_RE.findall(text)
    if not tokens:
        return

    table = doc.add_table(rows=1, cols=len(tokens))
    _remove_table_borders(table)
    _set_table_tight_spacing(table)

    for col, token in enumerate(tokens):
        cell = table.rows[0].cells[col]
        _set_cell_tight_margin(cell)
        _set_cell_valign(cell, "center")

        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = OVERLAY_LINE_SPACING
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if WORD_CHAR_RE.search(token):
            tr = translate_word_in_context(token, text, direction)
            tr_run = p.add_run(tr)
            set_run_font(tr_run, "Tahoma", OVERLAY_TR_SIZE)
            br_run = p.add_run()
            br_run.add_break()

        w_run = p.add_run(token)
        set_run_font(w_run, "Tahoma", OVERLAY_WORD_SIZE)

    row_h = OVERLAY_TR_SIZE + OVERLAY_WORD_SIZE + 1
    _set_row_exact_height(table.rows[0], row_h)

    tbl_element = table._tbl
    doc.element.body.remove(tbl_element)
    p_element = para._element
    p_element.addnext(tbl_element)
    p_element.getparent().remove(p_element)


def _apply_overlay_paragraph(para, direction: str, doc: Document):
    text = para.text
    if not text.strip():
        return

    if _paragraph_has_image(para):
        direction = resolve_direction(text, direction)
        rtl = direction == "en_ar" or is_mostly_arabic(text)
        set_paragraph_direction(para, rtl)
        for run in list(para.runs):
            if _run_has_image(run):
                continue
            _transform_text_run(para, run, direction)
        return

    _replace_paragraph_with_overlay_table(doc, para, direction)


def _process_docx_overlay(doc: Document, direction: str):
    for para in list(doc.paragraphs):
        _apply_overlay_paragraph(para, direction, doc)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in list(cell.paragraphs):
                    _apply_overlay_paragraph(para, direction, doc)


def _translate_pdf_overlay(source: Path, out_path: Path, direction: str):
    import fitz

    src = fitz.open(str(source))
    out = fitz.open()
    fontfile = find_arabic_font()

    for page_num in range(len(src)):
        page = src[page_num]
        new_page = out.new_page(width=page.rect.width, height=page.rect.height)
        new_page.show_pdf_page(page.rect, src, page_num)

        words = page.get_text("words")
        word_jobs = list(_iter_pdf_word_boxes(page))
        if not word_jobs and words:
            for w in words:
                token = w[4]
                if not WORD_CHAR_RE.search(token):
                    continue
                word_jobs.append((token, w[0], w[1], w[2], w[3], token))

        if not word_jobs:
            continue

        for token, x0, y0, x1, y1, line_text in word_jobs:
            tr = translate_word_in_context(token, line_text, direction)
            rtl = direction == "en_ar" or is_mostly_arabic(tr)
            _pdf_insert_translation_above(new_page, x0, y0, x1, y1, tr, fontfile, rtl)

    out.save(str(out_path))
    out.close()
    src.close()


def _build_overlay_file(
    source_path: Path, out_dir: Path, stem: str, direction: str, content: str
) -> dict[str, Path]:
    direction = resolve_direction(content, direction)
    suffix = source_path.suffix.lower()

    if suffix in (".docx", ".doc"):
        out = out_dir / f"{stem}_4_فوق_الكلمات.docx"
        shutil.copy2(source_path, out)
        doc = Document(out)
        _process_docx_overlay(doc, direction)
        doc.save(out)
        return {"overlay": out}

    if suffix == ".pdf":
        out = out_dir / f"{stem}_4_فوق_الكلمات.pdf"
        _translate_pdf_overlay(source_path, out, direction)
        return {"overlay": out}

    out = out_dir / f"{stem}_4_فوق_الكلمات.docx"
    doc = Document()
    for line in content.splitlines():
        if not line.strip():
            doc.add_paragraph()
            continue
        para = doc.add_paragraph(line)
        _replace_paragraph_with_overlay_table(doc, para, direction)
    doc.save(out)
    return {"overlay": out}


def _set_paragraph_translated(para, translated: str, direction: str):
    placed = False
    for run in para.runs:
        if run._element.xpath(".//w:drawing") or run._element.xpath(".//w:pict"):
            continue
        if not placed:
            run.text = translated
            set_run_font(run, "Tahoma", 12)
            placed = True
        else:
            run.text = ""
    if not placed:
        run = para.add_run(translated)
        set_run_font(run, "Tahoma", 12)
    style_paragraph(para, translated, direction)


def _process_docx_paragraphs(doc: Document, direction: str):
    for para in doc.paragraphs:
        if not para.text.strip():
            continue
        translated = translate_text(para.text.strip(), direction)
        _set_paragraph_translated(para, translated, direction)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if not para.text.strip():
                        continue
                    translated = translate_text(para.text.strip(), direction)
                    _set_paragraph_translated(para, translated, direction)


def _translate_docx(source: Path, out_dir: Path, stem: str, direction: str) -> dict[str, Path]:
    content = extract_text_from_file(source)
    _check_file_limits(source, content)
    direction = resolve_direction(content, direction)

    literal_files = _build_literal_files(content, direction, out_dir, stem)

    structured_path = out_dir / f"{stem}_2_بنفس_الترتيب.docx"
    shutil.copy2(source, structured_path)
    doc_full = Document(structured_path)
    _process_docx_paragraphs(doc_full, direction)
    doc_full.save(structured_path)

    return {**literal_files, "structured": structured_path}


def _pdf_write_in_box(page, rect, text: str, fontsize: float, fontfile: str | None, rtl: bool):
    import fitz
    from services.text_shape import has_arabic

    display = shape_for_pdf(text) if has_arabic(text) else text
    kwargs = {
        "fontsize": fontsize,
        "align": fitz.TEXT_ALIGN_RIGHT if rtl else fitz.TEXT_ALIGN_LEFT,
    }
    if fontfile:
        kwargs["fontfile"] = fontfile
        kwargs["fontname"] = "TahomaAr"
    try:
        page.insert_textbox(rect, display, **kwargs)
    except Exception:
        page.insert_textbox(rect, display, fontsize=fontsize, align=kwargs["align"])


def _build_structured_pdf(source: Path, out_path: Path, direction: str) -> None:
    import fitz

    src_doc = fitz.open(str(source))
    out_doc = fitz.open()
    fontfile = find_arabic_font()

    for page_num in range(len(src_doc)):
        page = src_doc[page_num]
        new_page = out_doc.new_page(width=page.rect.width, height=page.rect.height)
        new_page.show_pdf_page(page.rect, src_doc, page_num)

        redact_rects = []
        text_jobs = []

        for block in page.get_text("dict").get("blocks", []):
            if block.get("type") != 0:
                continue
            block_lines = []
            block_bbox = None
            block_size = 11
            for line in block.get("lines", []):
                line_text = "".join(s.get("text", "") for s in line.get("spans", []))
                if not line_text.strip():
                    continue
                block_lines.append(line_text.strip())
                if block_bbox is None:
                    block_bbox = fitz.Rect(line["bbox"])
                else:
                    block_bbox |= fitz.Rect(line["bbox"])
                block_size = max(block_size, line["spans"][0].get("size", 11))

            if not block_lines:
                continue
            full_text = " ".join(block_lines)
            new_text = translate_text(full_text, direction)
            redact_rects.append(block_bbox)
            rtl = direction == "en_ar" or is_mostly_arabic(new_text)
            text_jobs.append((block_bbox, new_text, block_size, rtl))

        for rect in redact_rects:
            new_page.add_redact_annot(rect, fill=(1, 1, 1))
        if redact_rects:
            new_page.apply_redactions()

        for rect, new_text, size, rtl in text_jobs:
            pad = fitz.Rect(rect.x0 - 2, rect.y0 - 2, rect.x1 + 40, rect.y1 + size * 3)
            _pdf_write_in_box(new_page, pad, new_text, size, fontfile, rtl)

    out_doc.save(str(out_path))
    out_doc.close()
    src_doc.close()


def prepare_full_file_translation(
    source_path: Path, output_dir: Path, direction: str = "auto",
) -> dict:
    output_dir.mkdir(parents=True, exist_ok=True)
    sample = extract_text_from_file(source_path)
    if not sample.strip():
        raise ValueError("لم يتم العثور على نص في الملف")
    _check_file_limits(source_path, sample)
    direction = resolve_direction(sample, direction)
    if not is_translator_ready():
        raise RuntimeError("محرك الترجمة غير جاهز — انتظر دقيقة ثم أعد المحاولة")
    _clear_word_cache()
    return {
        "source_path": source_path,
        "output_dir": output_dir,
        "content": sample,
        "direction": direction,
        "stem": source_path.stem,
        "suffix": source_path.suffix.lower(),
    }


def prepare_full_image_translation(
    image_path: Path, output_dir: Path, direction: str = "auto",
) -> dict:
    from services.ocr_service import ocr_image_layout

    output_dir.mkdir(parents=True, exist_ok=True)
    layout = ocr_image_layout(image_path)
    content = layout.get("text", "")
    if not content.strip():
        raise ValueError("لم يتم العثور على نص في الصورة")
    _check_file_limits(image_path, content)
    direction = resolve_direction(content, direction)
    if not is_translator_ready():
        raise RuntimeError("محرك الترجمة غير جاهز — انتظر دقيقة ثم أعد المحاولة")
    _clear_word_cache()
    return {
        "image_path": image_path,
        "layout": layout,
        "content": content,
        "direction": direction,
        "stem": image_path.stem,
        "output_dir": output_dir,
    }


def build_full_file_literal(data: dict) -> Path:
    set_file_translation_mode(True)
    return _build_literal_files(
        data["content"], data["direction"], data["output_dir"], data["stem"],
    )["literal"]


def build_full_file_structured(data: dict) -> Path:
    set_file_translation_mode(True)
    source = data["source_path"]
    out_dir = data["output_dir"]
    stem = data["stem"]
    direction = data["direction"]
    content = data["content"]
    suffix = data["suffix"]

    if suffix in (".docx", ".doc"):
        path = out_dir / f"{stem}_2_بنفس_الترتيب.docx"
        shutil.copy2(source, path)
        doc = Document(path)
        _process_docx_paragraphs(doc, direction)
        doc.save(path)
        return path
    if suffix == ".pdf":
        path = out_dir / f"{stem}_2_بنفس_الترتيب.pdf"
        _build_structured_pdf(source, path, direction)
        return path
    path = out_dir / f"{stem}_2_بنفس_الترتيب.txt"
    lines = [
        translate_text(u, direction)
        for u in _extract_logical_units(content) if u.strip()
    ]
    path.write_text("\n".join(lines), encoding="utf-8-sig")
    return path


def build_full_file_line_pairs(data: dict) -> Path:
    set_file_translation_mode(True)
    return _build_line_pairs_file(
        data["content"], data["direction"], data["output_dir"], data["stem"],
    )["line_pairs"]


def build_full_file_overlay(data: dict) -> Path:
    set_file_translation_mode(True)
    return _build_overlay_file(
        data["source_path"], data["output_dir"], data["stem"],
        data["direction"], data["content"],
    )["overlay"]


def build_full_image_literal(data: dict) -> Path:
    set_file_translation_mode(True)
    return _build_literal_files(
        data["content"], data["direction"], data["output_dir"], data["stem"],
    )["literal"]


def build_full_image_structured(data: dict) -> Path:
    set_file_translation_mode(True)
    path = data["output_dir"] / f"{data['stem']}_2_بنفس_الترتيب.pdf"
    _translate_image_structured(
        data["image_path"], data["layout"], path, data["direction"], data["content"],
    )
    return path


def build_full_image_line_pairs(data: dict) -> Path:
    set_file_translation_mode(True)
    return _build_line_pairs_file(
        data["content"], data["direction"], data["output_dir"], data["stem"],
    )["line_pairs"]


def build_full_image_overlay(data: dict) -> Path:
    set_file_translation_mode(True)
    path = data["output_dir"] / f"{data['stem']}_4_فوق_الكلمات.pdf"
    _translate_image_overlay(
        data["image_path"], data["layout"], path, data["direction"], data["content"],
    )
    return path


def _translate_pdf(source: Path, out_dir: Path, stem: str, direction: str) -> dict[str, Path]:
    content = extract_text_from_file(source)
    _check_file_limits(source, content)
    direction = resolve_direction(content, direction)

    literal_files = _build_literal_files(content, direction, out_dir, stem)
    outputs = dict(literal_files)
    out_path = out_dir / f"{stem}_2_بنفس_الترتيب.pdf"
    _build_structured_pdf(source, out_path, direction)
    outputs["structured"] = out_path
    return outputs


def _translate_txt(source: Path, out_dir: Path, stem: str, direction: str) -> dict[str, Path]:
    content = extract_text_from_file(source)
    _check_file_limits(source, content)
    direction = resolve_direction(content, direction)

    literal_files = _build_literal_files(content, direction, out_dir, stem)

    structured_lines = []
    for unit in _extract_logical_units(content):
        structured_lines.append(translate_text(unit, direction))

    structured_path = out_dir / f"{stem}_2_بنفس_الترتيب.txt"
    structured_path.write_text("\n".join(structured_lines), encoding="utf-8-sig")

    return {**literal_files, "structured": structured_path}


def _translate_image_structured(
    image_path: Path, layout: dict, out_path: Path, direction: str, content: str,
):
    import fitz

    direction = resolve_direction(content, direction)
    fontfile = find_arabic_font()
    doc = fitz.open()
    page = doc.new_page(width=layout["width"], height=layout["height"])
    page.insert_image(page.rect, filename=str(image_path))

    redact_rects = []
    text_jobs = []

    for block in layout.get("lines", []):
        line_text = block["text"].strip()
        if not line_text:
            continue
        rect = fitz.Rect(block["x0"], block["y0"], block["x1"], block["y1"])
        new_text = translate_text(line_text, direction)
        redact_rects.append(rect)
        block_h = max(block["y1"] - block["y0"], 8)
        rtl = direction == "en_ar" or is_mostly_arabic(new_text)
        text_jobs.append((rect, new_text, min(block_h, 14), rtl))

    for rect in redact_rects:
        page.add_redact_annot(rect, fill=(1, 1, 1))
    if redact_rects:
        page.apply_redactions()

    for rect, new_text, size, rtl in text_jobs:
        pad = fitz.Rect(rect.x0 - 2, rect.y0 - 2, rect.x1 + 40, rect.y1 + size * 2.5)
        _pdf_write_in_box(page, pad, new_text, size, fontfile, rtl)

    doc.save(str(out_path))
    doc.close()


def _translate_image_overlay(
    image_path: Path, layout: dict, out_path: Path, direction: str, content: str,
):
    import fitz

    direction = resolve_direction(content, direction)
    fontfile = find_arabic_font()
    doc = fitz.open()
    page = doc.new_page(width=layout["width"], height=layout["height"])
    page.insert_image(page.rect, filename=str(image_path))

    for token, x0, y0, x1, y1, line_text in layout.get("words", []):
        if not WORD_CHAR_RE.search(token):
            continue
        tr = translate_word_in_context(token, line_text, direction)
        rtl = direction == "en_ar" or is_mostly_arabic(tr)
        _pdf_insert_translation_above(page, x0, y0, x1, y1, tr, fontfile, rtl)

    doc.save(str(out_path))
    doc.close()


def _translate_paragraph_pairs(content: str, direction: str) -> list[tuple[str, str]]:
    """ترجمة فقرة بفقرة — سريعة ومناسبة للسحابة."""
    direction = resolve_direction(content, direction)
    pairs: list[tuple[str, str]] = []
    units = _extract_logical_units(content)
    total = len(units)
    for i, unit in enumerate(units, 1):
        unit = unit.strip()
        if not unit:
            continue
        if i % 10 == 0 or i == total:
            logger.info("Translating unit %d/%d", i, total)
        pairs.append((unit, translate_text(unit, direction)))
    return pairs


def _build_fast_outputs(
    pairs: list[tuple[str, str]], output_dir: Path, stem: str, direction: str,
) -> dict[str, Path]:
    """ملفان: PDF ثنائي + TXT كامل — بدون ترجمة كلمة بكلمة."""
    line_pdf = output_dir / f"{stem}_1_سطر_بسطر.pdf"
    create_pairs_pdf(
        pairs, line_pdf,
        title="ترجمة سطر بسطر — أصل وترجمة",
        direction=direction,
    )

    full_txt = output_dir / f"{stem}_2_ترجمة_كاملة.txt"
    full_txt.write_text("\n\n".join(tr for _, tr in pairs), encoding="utf-8-sig")

    return {"line_pairs": line_pdf, "structured": full_txt}


def translate_file_fast(
    source_path: Path, output_dir: Path, direction: str = "auto",
) -> dict[str, Path]:
    _clear_word_cache()
    output_dir.mkdir(parents=True, exist_ok=True)
    stem = source_path.stem

    content = extract_text_from_file(source_path)
    if not content.strip():
        raise ValueError("لم يتم العثور على نص في الملف")
    _check_file_limits(source_path, content)
    direction = resolve_direction(content, direction)

    logger.info("Fast file translation: %s (%d chars)", source_path.name, len(content))
    pairs = _translate_paragraph_pairs(content, direction)
    if not pairs:
        raise ValueError("لم يتم العثور على فقرات للترجمة")
    return _build_fast_outputs(pairs, output_dir, stem, direction)


def translate_image_fast(
    image_path: Path, output_dir: Path, direction: str = "auto",
) -> dict[str, Path]:
    _clear_word_cache()
    from services.ocr_service import ocr_image

    output_dir.mkdir(parents=True, exist_ok=True)
    stem = image_path.stem

    content = ocr_image(image_path)
    if not content.strip():
        raise ValueError("لم يتم العثور على نص في الصورة")
    _check_file_limits(image_path, content)
    direction = resolve_direction(content, direction)

    logger.info("Fast image translation: %s (%d chars)", image_path.name, len(content))
    pairs = _translate_paragraph_pairs(content, direction)
    if not pairs:
        raise ValueError("لم يتم العثور على نص للترجمة في الصورة")
    return _build_fast_outputs(pairs, output_dir, stem, direction)


def prepare_dual_file_translation(
    source_path: Path, output_dir: Path, direction: str = "auto",
) -> dict:
    data = prepare_full_file_translation(source_path, output_dir, direction)
    set_file_translation_mode(True)
    logger.info("Prewarming Argos cache for fast overlay...")
    _prewarm_word_cache_for_text(data["content"], data["direction"])
    return data


def prepare_dual_image_translation(
    image_path: Path, output_dir: Path, direction: str = "auto",
) -> dict:
    data = prepare_full_image_translation(image_path, output_dir, direction)
    set_file_translation_mode(True)
    logger.info("Prewarming Argos cache for fast overlay...")
    _prewarm_word_cache_for_text(data["content"], data["direction"])
    return data


def translate_file_dual_modes(
    source_path: Path, output_dir: Path, direction: str = "auto",
) -> dict[str, Path]:
    """ملفان فقط: بنفس الترتيب + فوق الكلمات — Argos محلي سريع."""
    set_file_translation_mode(True)
    try:
        data = prepare_dual_file_translation(source_path, output_dir, direction)
        logger.info("Dual file translation (2 files): %s", source_path.name)
        structured = build_full_file_structured(data)
        overlay = build_full_file_overlay(data)
        return {"structured": structured, "overlay": overlay}
    finally:
        set_file_translation_mode(False)


def translate_image_dual_modes(
    image_path: Path, output_dir: Path, direction: str = "auto",
) -> dict[str, Path]:
    set_file_translation_mode(True)
    try:
        data = prepare_dual_image_translation(image_path, output_dir, direction)
        logger.info("Dual image translation (2 files): %s", image_path.name)
        structured = build_full_image_structured(data)
        overlay = build_full_image_overlay(data)
        return {"structured": structured, "overlay": overlay}
    finally:
        set_file_translation_mode(False)


def translate_image_two_modes(
    image_path: Path, output_dir: Path, direction: str = "auto"
) -> dict[str, Path]:
    if use_fast_file_translation():
        return translate_image_fast(image_path, output_dir, direction)
    if use_full_file_translation():
        set_file_translation_mode(True)
        try:
            return _translate_image_full(image_path, output_dir, direction)
        finally:
            set_file_translation_mode(False)
    return translate_image_dual_modes(image_path, output_dir, direction)


def _translate_image_full(
    image_path: Path, output_dir: Path, direction: str = "auto",
) -> dict[str, Path]:
    _clear_word_cache()
    from services.ocr_service import ocr_image_layout

    output_dir.mkdir(parents=True, exist_ok=True)
    stem = image_path.stem
    layout = ocr_image_layout(image_path)
    content = layout.get("text", "")
    if not content.strip():
        raise ValueError("لم يتم العثور على نص في الصورة")
    _check_file_limits(image_path, content)
    direction = resolve_direction(content, direction)
    logger.info("Full image translation (4 files): %s", image_path.name)
    logger.info("Step 1/4: literal PDF...")
    result = _build_literal_files(content, direction, output_dir, stem)

    logger.info("Step 2/4: structured PDF...")
    structured_pdf = output_dir / f"{stem}_2_بنفس_الترتيب.pdf"
    _translate_image_structured(image_path, layout, structured_pdf, direction, content)
    result["structured"] = structured_pdf

    logger.info("Step 3/4: line pairs...")
    result.update(_build_line_pairs_file(content, direction, output_dir, stem))

    logger.info("Step 4/4: overlay...")
    overlay_pdf = output_dir / f"{stem}_4_فوق_الكلمات.pdf"
    _translate_image_overlay(image_path, layout, overlay_pdf, direction, content)
    result["overlay"] = overlay_pdf

    logger.info("Full image translation done")
    return result


def translate_file_two_modes(
    source_path: Path, output_dir: Path, direction: str = "auto"
) -> dict[str, Path]:
    if use_fast_file_translation():
        return translate_file_fast(source_path, output_dir, direction)
    if use_full_file_translation():
        set_file_translation_mode(True)
        try:
            return _translate_file_full(source_path, output_dir, direction)
        finally:
            set_file_translation_mode(False)
    return translate_file_dual_modes(source_path, output_dir, direction)


def _translate_file_full(
    source_path: Path, output_dir: Path, direction: str = "auto",
) -> dict[str, Path]:
    _clear_word_cache()
    suffix = source_path.suffix.lower()
    stem = source_path.stem
    output_dir.mkdir(parents=True, exist_ok=True)

    sample = extract_text_from_file(source_path)
    if not sample.strip():
        raise ValueError("لم يتم العثور على نص في الملف")
    direction = resolve_direction(sample, direction)
    logger.info("Full translation (4 files): %s", source_path.name)
    logger.info("Step 1-2/4: literal + structured...")
    if suffix in (".docx", ".doc"):
        result = _translate_docx(source_path, output_dir, stem, direction)
    elif suffix == ".pdf":
        result = _translate_pdf(source_path, output_dir, stem, direction)
    elif suffix == ".txt":
        result = _translate_txt(source_path, output_dir, stem, direction)
    else:
        raise ValueError(f"نوع الملف غير مدعوم: {suffix}")

    logger.info("Step 3/4: line pairs...")
    result.update(_build_line_pairs_file(sample, direction, output_dir, stem))
    logger.info("Step 4/4: overlay...")
    result.update(_build_overlay_file(source_path, output_dir, stem, direction, sample))
    logger.info("Full translation done: %s", list(result.keys()))
    return result


# توافق مع الكود القديم
translate_file_three_modes = translate_file_two_modes
